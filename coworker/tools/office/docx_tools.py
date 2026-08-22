"""Word documents — read, write, and edit .docx deliverables.

The model cannot emit valid OOXML, but it emits structured JSON reliably, so writing goes
through a small document IR (``heading`` / ``paragraph`` / ``bullet`` / ``table`` /
``page_break``) rather than raw markup.

Reading mirrors ``tools/files.py``: blocks are **numbered** so the model can cite "block 12"
and then edit exactly that block. ``edit_document`` mutates paragraphs in place through
python-docx, so styles, headers, numbering, and everything the model never saw survive the
edit — the failure mode of "rewrite the whole file" is losing all of it.
"""

from __future__ import annotations

from typing import Any

from ... import deliverable_check
from ._common import MAX_TEXT_CHARS, clip, decorate, guard, require
from .paths import context_roots, display_path, resolve_read, resolve_write

_DEFAULT_LIMIT = 200

_BLOCK_SHAPE = {
    "type": "object",
    "properties": {
        "type": {
            "type": "string",
            "enum": ["heading", "paragraph", "bullet", "numbered", "table", "page_break"],
            "description": "Block kind.",
        },
        "text": {"type": "string", "description": "Text for non-table blocks."},
        "level": {
            "type": "integer",
            "description": "Heading level 1-9 (heading blocks only; default 1).",
        },
        "rows": {
            "type": "array",
            "description": "Table blocks only: a list of rows, each a list of cell strings.",
            "items": {"type": "array", "items": {"type": "string"}},
        },
        "header": {
            "type": "boolean",
            "description": "Table blocks only: style the first row as a header (default true).",
        },
    },
    "required": ["type"],
}

_WRITE_SCHEMA = {
    "type": "function",
    "function": {
        "name": "write_document",
        "description": (
            "Create or overwrite a Word (.docx) document from structured blocks. Use this for "
            "any Word deliverable — do NOT write a script to do it. Pass append=true to add "
            "blocks to an existing document instead of replacing it."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {
                    "type": "string",
                    "description": "Destination .docx path, relative to the workspace.",
                },
                "blocks": {
                    "type": "array",
                    "description": "Document content, in order.",
                    "items": _BLOCK_SHAPE,
                },
                "append": {
                    "type": "boolean",
                    "description": "Append to an existing document (default false = overwrite).",
                },
            },
            "required": ["path", "blocks"],
        },
    },
}

_READ_SCHEMA = {
    "type": "function",
    "function": {
        "name": "read_document",
        "description": (
            "Read a Word (.docx) document as numbered blocks (headings, paragraphs, bullets, "
            "tables). Use the returned block index with edit_document. Long documents are "
            "windowed: pass start to continue where the previous read stopped. Read-only."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "The .docx file to read."},
                "start": {
                    "type": "integer",
                    "description": "First block index to return, 0-based (default 0).",
                },
                "limit": {
                    "type": "integer",
                    "description": f"How many blocks (default {_DEFAULT_LIMIT}).",
                },
                "revisions": {
                    "type": "boolean",
                    "description": (
                        "Also list pending tracked changes (author, deleted/inserted text) "
                        "under `revisions`."
                    ),
                },
            },
            "required": ["path"],
        },
    },
}

_REVISE_SCHEMA = {
    "type": "function",
    "function": {
        "name": "revise_document",
        "description": (
            "Revise text blocks of an existing Word document as TRACKED CHANGES (Word's "
            "Review ▸ Track Changes): the old text stays as a deletion and the new text is an "
            "insertion, both attributed to Mimi, so the user accepts or rejects each change in "
            "Word. Use this for documents the user wrote or that others will review; use "
            "edit_document only for direct, silent edits to your own drafts. Read the document "
            "first to get block indexes. The result lists every change as before → after with "
            "your reason — repeat that list to the user in plain language."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "The .docx file to revise."},
                "edits": {
                    "type": "array",
                    "description": "Revisions to apply.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "index": {
                                "type": "integer",
                                "description": "Block index from read_document.",
                            },
                            "text": {
                                "type": "string",
                                "description": "The replacement text for the whole block.",
                            },
                            "reason": {
                                "type": "string",
                                "description": (
                                    "One plain-language sentence on why (e.g. 'tightened the "
                                    "claim to what the data supports')."
                                ),
                            },
                        },
                        "required": ["index", "text"],
                    },
                },
            },
            "required": ["path", "edits"],
        },
    },
}

_EDIT_SCHEMA = {
    "type": "function",
    "function": {
        "name": "edit_document",
        "description": (
            "Replace the text of specific blocks in an existing Word document, in place. "
            "Styles, headers, and the rest of the document are preserved. Read the document "
            "first to get block indexes."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "path": {"type": "string", "description": "The .docx file to edit."},
                "edits": {
                    "type": "array",
                    "description": "Replacements to apply.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "index": {
                                "type": "integer",
                                "description": "Block index from read_document.",
                            },
                            "text": {"type": "string", "description": "Replacement text."},
                        },
                        "required": ["index", "text"],
                    },
                },
            },
            "required": ["path", "edits"],
        },
    },
}


def _style_of(paragraph: Any) -> tuple[str, int]:
    """Map a python-docx paragraph style back onto the IR (kind, heading level)."""
    name = (getattr(paragraph.style, "name", "") or "").strip()
    if name.startswith("Heading"):
        tail = name.replace("Heading", "").strip()
        return "heading", int(tail) if tail.isdigit() else 1
    if name.startswith("List Bullet"):
        return "bullet", 0
    if name.startswith("List Number"):
        return "numbered", 0
    return "paragraph", 0


def _add_block(document: Any, block: Any) -> None:
    if not isinstance(block, dict):
        raise ValueError(f"each block must be an object, got {type(block).__name__}")
    kind = str(block.get("type") or "paragraph").lower()
    text = str(block.get("text") or "")

    if kind == "heading":
        level = block.get("level", 1)
        level = level if isinstance(level, int) and 1 <= level <= 9 else 1
        document.add_heading(text, level=level)
    elif kind == "paragraph":
        document.add_paragraph(text)
    elif kind == "bullet":
        document.add_paragraph(text, style="List Bullet")
    elif kind == "numbered":
        document.add_paragraph(text, style="List Number")
    elif kind == "page_break":
        document.add_page_break()
    elif kind == "table":
        rows = block.get("rows") or []
        if not rows:
            raise ValueError("a table block needs a non-empty 'rows' list")
        width = max(len(r) for r in rows)
        table = document.add_table(rows=len(rows), cols=width)
        table.style = "Table Grid"
        for r, row in enumerate(rows):
            for c in range(width):
                value = row[c] if c < len(row) else ""
                table.cell(r, c).text = "" if value is None else str(value)
        if block.get("header", True) and rows:
            for cell in table.rows[0].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True
    else:
        raise ValueError(
            f"unknown block type {kind!r}; expected one of: heading, paragraph, bullet, "
            "numbered, table, page_break"
        )


# -- tracked changes (WordprocessingML revisions) ---------------------------------------------
# A replacement becomes <w:del> around the paragraph's existing runs (their w:t → w:delText)
# followed by one <w:ins> run carrying the first run's formatting. Word, LibreOffice and Pages
# all render these as review marks with accept/reject. Ids must be unique per document.
_REV_AUTHOR = "Mimi"


def _revision_stamp() -> str:
    import datetime as _dt

    return _dt.datetime.now(_dt.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _para_text(para: Any) -> str:
    """The paragraph's ACCEPTED text: every w:t under it (runs inside <w:ins> included),
    never w:delText. python-docx's `.text` reads direct runs only, so a revised paragraph
    would otherwise read as empty and vanish from the block index."""
    from docx.oxml.ns import qn

    return "".join((t.text or "") for t in para._element.iter(qn("w:t")))


def _indexed_blocks(document: Any) -> list[Any]:
    """The block index space read_document exposes: non-empty paragraphs + tables (None)."""
    indexed: list[Any] = []
    paragraphs = {p._element: p for p in document.paragraphs}
    tables = {t._element: t for t in document.tables}
    for child in document.element.body.iterchildren():
        if child in paragraphs:
            para = paragraphs[child]
            if _para_text(para).strip():
                indexed.append(para)
        elif child in tables:
            indexed.append(None)
    return indexed


def _next_revision_id(document: Any) -> int:
    from docx.oxml.ns import qn

    highest = 0
    for tag in ("w:ins", "w:del"):
        for el in document.element.body.iter(qn(tag)):
            try:
                highest = max(highest, int(el.get(qn("w:id")) or 0))
            except ValueError:
                continue
    return highest + 1


def _track_replacement(para: Any, new_text: str, *, rev_id: int, stamp: str) -> int:
    """Wrap the paragraph's runs in <w:del>, append <w:ins> with the new text. Returns the
    next free revision id."""
    import copy

    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = para._element
    # Live runs = direct children plus runs inside earlier <w:ins> (a second revision of the
    # same paragraph deletes the previously inserted text); runs already under <w:del> stay.
    runs = [
        r
        for r in p.iter(qn("w:r"))
        if r.getparent() is p or r.getparent().tag == qn("w:ins")
    ]
    rpr = None
    if runs:
        first_rpr = runs[0].find(qn("w:rPr"))
        rpr = copy.deepcopy(first_rpr) if first_rpr is not None else None

    if runs:
        deletion = OxmlElement("w:del")
        deletion.set(qn("w:id"), str(rev_id))
        deletion.set(qn("w:author"), _REV_AUTHOR)
        deletion.set(qn("w:date"), stamp)
        rev_id += 1
        runs[0].addprevious(deletion)
        for run in runs:
            for t in run.findall(qn("w:t")):
                t.tag = qn("w:delText")
            deletion.append(run)  # moves the run under <w:del>

    insertion = OxmlElement("w:ins")
    insertion.set(qn("w:id"), str(rev_id))
    insertion.set(qn("w:author"), _REV_AUTHOR)
    insertion.set(qn("w:date"), stamp)
    rev_id += 1
    new_run = OxmlElement("w:r")
    if rpr is not None:
        new_run.append(rpr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    new_run.append(t)
    insertion.append(new_run)
    p.append(insertion)
    return rev_id


def _list_revisions(document: Any) -> list[dict[str, Any]]:
    """Pending tracked changes, in read_document's index space: {index, author, deleted, inserted}."""
    from docx.oxml.ns import qn

    out: list[dict[str, Any]] = []
    for index, para in enumerate(_indexed_blocks(document)):
        if para is None:
            continue
        p = para._element
        dels = p.findall(qn("w:del"))
        inss = p.findall(qn("w:ins"))
        if not dels and not inss:
            continue
        deleted = "".join((t.text or "") for d in dels for t in d.iter(qn("w:delText")))
        inserted = "".join((t.text or "") for i in inss for t in i.iter(qn("w:t")))
        author = next(
            (el.get(qn("w:author")) for el in (dels + inss) if el.get(qn("w:author"))),
            "",
        )
        out.append(
            {
                "index": index,
                "author": author,
                "deleted": clip(deleted, 400),
                "inserted": clip(inserted, 400),
            }
        )
    return out


def docx_tools(context: Any) -> list:
    roots = context_roots(context)

    @guard
    def write_document(path: str, blocks: list, append: bool = False) -> dict[str, Any]:
        docx = require("docx", "python-docx")
        target = resolve_write(path, roots)
        if not isinstance(blocks, list):
            raise ValueError("'blocks' must be a list of block objects")

        if append and target.is_file():
            document = docx.Document(str(target))
        else:
            document = docx.Document()
        for block in blocks:
            _add_block(document, block)

        target.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(target))
        return deliverable_check.attach({
            "path": display_path(target, roots),
            "blocks_written": len(blocks),
            "appended": bool(append),
            "bytes": target.stat().st_size,
        }, target)

    @guard
    def read_document(
        path: str, start: int = 0, limit: int = _DEFAULT_LIMIT, revisions: bool = False
    ) -> dict[str, Any]:
        docx = require("docx", "python-docx")
        target = resolve_read(path, roots)
        if not target.is_file():
            raise FileNotFoundError(display_path(target, roots))

        document = docx.Document(str(target))
        # Walk the body in document order so a table between two paragraphs keeps its place;
        # python-docx's .paragraphs / .tables are separate sequences and lose the interleaving.
        blocks: list[dict[str, Any]] = []
        body = document.element.body
        paragraphs = {p._element: p for p in document.paragraphs}
        tables = {t._element: t for t in document.tables}
        for child in body.iterchildren():
            if child in paragraphs:
                para = paragraphs[child]
                text = _para_text(para).strip()  # accepted view of any tracked changes
                if not text:
                    continue
                kind, level = _style_of(para)
                entry: dict[str, Any] = {"type": kind, "text": clip(text, MAX_TEXT_CHARS)}
                if kind == "heading":
                    entry["level"] = level
                blocks.append(entry)
            elif child in tables:
                table = tables[child]
                blocks.append(
                    {
                        "type": "table",
                        "rows": [
                            [clip(cell.text.strip()) for cell in row.cells]
                            for row in table.rows
                        ],
                    }
                )

        total = len(blocks)
        begin = start if isinstance(start, int) and start > 0 else 0
        count = limit if isinstance(limit, int) and limit > 0 else _DEFAULT_LIMIT
        window = blocks[begin : begin + count]
        for offset, block in enumerate(window):
            block["index"] = begin + offset

        result: dict[str, Any] = {
            "path": display_path(target, roots),
            "total_blocks": total,
            "blocks": window,
        }
        if revisions:
            result["revisions"] = _list_revisions(document)
        end = begin + len(window)
        if end < total:
            result["note"] = (
                f"showing blocks {begin}-{end - 1} of {total}; "
                f"call again with start={end} to continue"
            )
        return result

    @guard
    def edit_document(path: str, edits: list) -> dict[str, Any]:
        docx = require("docx", "python-docx")
        target = resolve_write(path, roots)
        if not target.is_file():
            raise FileNotFoundError(display_path(target, roots))
        if not isinstance(edits, list) or not edits:
            raise ValueError("'edits' must be a non-empty list")

        document = docx.Document(str(target))
        # Rebuild the same index space read_document exposes: non-empty paragraphs and tables,
        # in document order. Only paragraphs are text-editable.
        indexed: list[Any] = []
        paragraphs = {p._element: p for p in document.paragraphs}
        tables = {t._element: t for t in document.tables}
        for child in document.element.body.iterchildren():
            if child in paragraphs:
                para = paragraphs[child]
                if para.text.strip():
                    indexed.append(para)
            elif child in tables:
                indexed.append(None)  # a table occupies an index but is not text-editable

        applied = 0
        for edit in edits:
            if not isinstance(edit, dict):
                raise ValueError("each edit must be an object with 'index' and 'text'")
            index = edit.get("index")
            if not isinstance(index, int) or not 0 <= index < len(indexed):
                raise ValueError(
                    f"block index {index!r} is out of range (document has {len(indexed)} blocks)"
                )
            para = indexed[index]
            if para is None:
                raise ValueError(f"block {index} is a table; edit_document only edits text blocks")
            # Write into the first run so its formatting (font, bold, size) is kept, and drop
            # the rest — replacing paragraph.text outright would discard the run styling.
            runs = para.runs
            if runs:
                runs[0].text = str(edit.get("text") or "")
                for extra in runs[1:]:
                    extra.text = ""
            else:
                para.text = str(edit.get("text") or "")
            applied += 1

        document.save(str(target))
        return {"path": display_path(target, roots), "edited": applied}

    @guard
    def revise_document(path: str, edits: list) -> dict[str, Any]:
        docx = require("docx", "python-docx")
        target = resolve_write(path, roots)
        if not target.is_file():
            raise FileNotFoundError(display_path(target, roots))
        if not isinstance(edits, list) or not edits:
            raise ValueError("'edits' must be a non-empty list")

        document = docx.Document(str(target))
        indexed = _indexed_blocks(document)
        next_id = _next_revision_id(document)
        stamp = _revision_stamp()
        changes: list[dict[str, Any]] = []
        for edit in edits:
            if not isinstance(edit, dict):
                raise ValueError("each edit must be an object with 'index' and 'text'")
            index = edit.get("index")
            if not isinstance(index, int) or not 0 <= index < len(indexed):
                raise ValueError(
                    f"block index {index!r} is out of range (document has {len(indexed)} blocks)"
                )
            para = indexed[index]
            if para is None:
                raise ValueError(f"block {index} is a table; revise_document only revises text blocks")
            before = _para_text(para)
            after = str(edit.get("text") or "")
            if before == after:
                continue  # nothing to track
            next_id = _track_replacement(para, after, rev_id=next_id, stamp=stamp)
            changes.append(
                {
                    "index": index,
                    "before": clip(before, 400),
                    "after": clip(after, 400),
                    "reason": str(edit.get("reason") or "").strip(),
                }
            )

        document.save(str(target))
        return {
            "path": display_path(target, roots),
            "applied": len(changes),
            "changes": changes,
            "note": (
                "Changes are tracked — the user accepts or rejects them in Word "
                "(Review ▸ Track Changes). Tell the user what changed and why, in plain language."
            ),
        }

    return [
        decorate(
            write_document,
            name="write_document",
            schema=_WRITE_SCHEMA,
            risk="medium",
            capabilities=["write"],
        ),
        decorate(read_document, name="read_document", schema=_READ_SCHEMA),
        decorate(
            edit_document,
            name="edit_document",
            schema=_EDIT_SCHEMA,
            risk="medium",
            capabilities=["write"],
        ),
        decorate(
            revise_document,
            name="revise_document",
            schema=_REVISE_SCHEMA,
            risk="medium",
            capabilities=["write"],
        ),
    ]
